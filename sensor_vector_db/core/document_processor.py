"""Document parsing, OCR fallback, metadata extraction, and chunking."""

from __future__ import annotations

from abc import ABC, abstractmethod
import ast
import json
import os
from pathlib import Path
import re
import tempfile
from typing import Any, Callable

from sensor_vector_db.config.settings import Settings, get_settings
from sensor_vector_db.core.types import OperationCancelled, ParsedSegment, TextChunk
from sensor_vector_db.utils.file_utils import (
    detect_file_type,
    get_file_info,
    read_text_with_fallback,
)
from sensor_vector_db.utils.logger import get_logger


logger = get_logger(__name__)
CancelCallback = Callable[[], None]


class BaseDocumentParser(ABC):
    """Abstract parser interface for supported document formats."""

    def __init__(self, settings: Settings | None = None) -> None:
        """Initialize a parser with runtime settings."""
        self.settings = settings or get_settings()

    @abstractmethod
    def parse(
        self,
        path: str | Path,
        cancel_callback: CancelCallback | None = None,
    ) -> list[ParsedSegment]:
        """Parse a source file into source-preserving segments."""


class OCRClient:
    """Lazy PaddleOCR wrapper used only for scanned PDF pages."""

    def __init__(self, settings: Settings | None = None) -> None:
        """Create an OCR client without loading PaddleOCR immediately."""
        self.settings = settings or get_settings()
        self._ocr: Any | None = None

    def _load(self) -> Any:
        """Load PaddleOCR on demand."""
        if self._ocr is not None:
            return self._ocr
        os.environ.setdefault("FLAGS_use_mkldnn", "0")
        os.environ.setdefault("FLAGS_use_onednn", "0")
        os.environ.setdefault("PADDLE_PDX_ENABLE_MKLDNN_BYDEFAULT", "0")
        os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
        try:
            from paddleocr import PaddleOCR
        except ImportError as exc:
            raise RuntimeError(
                "PaddleOCR is not installed. Install requirements.txt to enable OCR."
            ) from exc

        try:
            self._ocr = PaddleOCR(lang=self.settings.ocr_lang)
        except TypeError:
            self._ocr = PaddleOCR(use_angle_cls=True, lang=self.settings.ocr_lang)
        return self._ocr

    def image_to_text(self, image_path: str | Path) -> str:
        """Run OCR on an image and return recognized text."""
        ocr = self._load()
        try:
            result = ocr.predict(str(image_path))
        except AttributeError:
            result = ocr.ocr(str(image_path), cls=True)
        except Exception as exc:
            raise RuntimeError(f"OCR failed for {image_path}: {exc}") from exc
        return self._flatten_result(result)

    @staticmethod
    def _flatten_result(result: Any) -> str:
        """Extract text strings from common PaddleOCR result shapes."""
        texts: list[str] = []

        def walk(value: Any) -> None:
            if value is None:
                return
            if isinstance(value, dict):
                for key in ("rec_texts", "text", "texts"):
                    item = value.get(key)
                    if isinstance(item, str):
                        texts.append(item)
                    elif isinstance(item, list):
                        texts.extend(str(part) for part in item if str(part).strip())
                for item in value.values():
                    walk(item)
                return
            if isinstance(value, (list, tuple)):
                if len(value) >= 2 and isinstance(value[1], (list, tuple)):
                    maybe_text = value[1][0] if value[1] else None
                    if isinstance(maybe_text, str):
                        texts.append(maybe_text)
                for item in value:
                    walk(item)

        walk(result)
        deduped = []
        seen = set()
        for text in texts:
            cleaned = str(text).strip()
            if cleaned and cleaned not in seen:
                deduped.append(cleaned)
                seen.add(cleaned)
        return "\n".join(deduped)


class PDFParser(BaseDocumentParser):
    """Parse PDF text and tables with optional OCR fallback."""

    def __init__(self, settings: Settings | None = None) -> None:
        """Initialize PDF parser."""
        super().__init__(settings)
        self.ocr_client = OCRClient(self.settings)

    def parse(
        self,
        path: str | Path,
        cancel_callback: CancelCallback | None = None,
    ) -> list[ParsedSegment]:
        """Parse PDF pages into text, table, and OCR segments."""
        _check_cancelled(cancel_callback)
        file_path = Path(path)
        segments: list[ParsedSegment] = []
        try:
            import pdfplumber
        except ImportError as exc:
            raise RuntimeError("pdfplumber is required to parse PDF files.") from exc

        try:
            with pdfplumber.open(file_path) as pdf:
                pdfium_document: Any | None = None
                ocr_pages_used = 0
                try:
                    for page_number, page in enumerate(pdf.pages, start=1):
                        _check_cancelled(cancel_callback)
                        text = page.extract_text() or ""
                        if text.strip():
                            segments.append(
                                ParsedSegment(
                                    content=text.strip(),
                                    content_type="text",
                                    page_number=page_number,
                                )
                            )

                        for table in page.extract_tables() or []:
                            table_text = _table_to_markdown(table)
                            if table_text.strip():
                                segments.append(
                                    ParsedSegment(
                                        content=table_text,
                                        content_type="table",
                                        page_number=page_number,
                                    )
                                )

                        if self._should_ocr_page(text, ocr_pages_used):
                            _check_cancelled(cancel_callback)
                            if pdfium_document is None:
                                pdfium_document = self._open_pdfium_document(file_path)
                            ocr_pages_used += 1
                            ocr_text = self._ocr_pdf_page(pdfium_document, file_path, page_number - 1)
                            _check_cancelled(cancel_callback)
                            if ocr_text.strip():
                                segments.append(
                                    ParsedSegment(
                                        content=ocr_text.strip(),
                                        content_type="ocr",
                                        page_number=page_number,
                                    )
                                )
                finally:
                    self._close_pdfium_document(pdfium_document)
            return segments
        except OperationCancelled:
            raise
        except Exception as exc:
            raise RuntimeError(f"Failed to parse PDF {file_path}: {exc}") from exc

    def _should_ocr_page(self, text: str, ocr_pages_used: int = 0) -> bool:
        """Return whether a PDF page should be sent to OCR."""
        if not self.settings.ocr_enabled:
            return False
        if len(text.strip()) >= self.settings.ocr_min_text_chars:
            return False
        max_pages = self.settings.ocr_max_pages_per_file
        return max_pages <= 0 or ocr_pages_used < max_pages

    @staticmethod
    def _open_pdfium_document(path: Path) -> Any:
        """Open a PDF with pypdfium2 for OCR rendering."""
        try:
            import pypdfium2 as pdfium
        except ImportError as exc:
            raise RuntimeError("pypdfium2 is required for PDF OCR rendering.") from exc
        return pdfium.PdfDocument(str(path))

    @staticmethod
    def _close_pdfium_document(pdfium_document: Any | None) -> None:
        """Close a pypdfium2 document when the runtime exposes a close method."""
        if pdfium_document is None:
            return
        close = getattr(pdfium_document, "close", None)
        if callable(close):
            close()

    def _ocr_pdf_page(self, pdfium_document: Any, path: Path, page_index: int) -> str:
        """Render a PDF page to a temporary image and run OCR."""
        try:
            page = pdfium_document[page_index]
            bitmap = page.render(scale=self.settings.ocr_render_scale)
            pil_image = bitmap.to_pil()
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
                temp_path = Path(temp_file.name)
            pil_image.save(temp_path)
            try:
                return self.ocr_client.image_to_text(temp_path)
            finally:
                temp_path.unlink(missing_ok=True)
        except Exception as exc:
            logger.warning("OCR fallback failed for %s page %s: %s", path, page_index + 1, exc)
            return ""


class WordParser(BaseDocumentParser):
    """Parse DOCX paragraphs and tables."""

    def parse(
        self,
        path: str | Path,
        cancel_callback: CancelCallback | None = None,
    ) -> list[ParsedSegment]:
        """Parse a DOCX document."""
        _check_cancelled(cancel_callback)
        file_path = Path(path)
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("python-docx is required to parse DOCX files.") from exc

        try:
            document = DocxDocument(str(file_path))
            segments: list[ParsedSegment] = []
            paragraph_text = "\n".join(
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            )
            if paragraph_text:
                segments.append(ParsedSegment(paragraph_text, "text"))
            for table in document.tables:
                _check_cancelled(cancel_callback)
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                table_text = _table_to_markdown(rows)
                if table_text.strip():
                    segments.append(ParsedSegment(table_text, "table"))
            return segments
        except OperationCancelled:
            raise
        except Exception as exc:
            raise RuntimeError(f"Failed to parse DOCX {file_path}: {exc}") from exc


class TextParser(BaseDocumentParser):
    """Parse plain text-like files."""

    def parse(
        self,
        path: str | Path,
        cancel_callback: CancelCallback | None = None,
    ) -> list[ParsedSegment]:
        """Read a text file as one segment."""
        _check_cancelled(cancel_callback)
        content = read_text_with_fallback(path)
        return [ParsedSegment(content=content, content_type="text")]


class CodeParser(BaseDocumentParser):
    """Parse source code files while preserving syntax units where possible."""

    def parse(
        self,
        path: str | Path,
        cancel_callback: CancelCallback | None = None,
    ) -> list[ParsedSegment]:
        """Parse code by AST for Python or conservative line chunks otherwise."""
        _check_cancelled(cancel_callback)
        file_path = Path(path)
        content = read_text_with_fallback(file_path)
        if file_path.suffix.lower() == ".py":
            python_segments = self._parse_python(content)
            if python_segments:
                return python_segments
        return [ParsedSegment(content=content, content_type="code")]

    def _parse_python(self, content: str) -> list[ParsedSegment]:
        """Split Python source into classes/functions when AST line numbers exist."""
        try:
            tree = ast.parse(content)
        except SyntaxError:
            return []

        lines = content.splitlines()
        segments: list[ParsedSegment] = []
        for node in tree.body:
            if isinstance(node, (ast.ClassDef, ast.FunctionDef, ast.AsyncFunctionDef)):
                start = getattr(node, "lineno", 1)
                end = getattr(node, "end_lineno", start)
                snippet = "\n".join(lines[start - 1 : end])
                if snippet.strip():
                    segments.append(
                        ParsedSegment(
                            content=snippet,
                            content_type="code",
                            metadata={"symbol": node.name, "start_line": start, "end_line": end},
                        )
                    )
        return segments


class DocumentParserFactory:
    """Factory selecting the parser for a supported file."""

    def __init__(self, settings: Settings | None = None) -> None:
        """Initialize factory."""
        self.settings = settings or get_settings()
        self._parsers: dict[str, BaseDocumentParser] = {}

    def get_parser(self, path: str | Path) -> BaseDocumentParser:
        """Return a parser for the file type."""
        file_type = detect_file_type(path)
        parser = self._parsers.get(file_type)
        if parser:
            return parser
        if file_type == "pdf":
            parser = PDFParser(self.settings)
        elif file_type == "docx":
            parser = WordParser(self.settings)
        elif file_type == "text":
            parser = TextParser(self.settings)
        elif file_type == "code":
            parser = CodeParser(self.settings)
        else:
            raise ValueError(f"Unsupported file type for {path}")
        self._parsers[file_type] = parser
        return parser

    def parse(
        self,
        path: str | Path,
        cancel_callback: CancelCallback | None = None,
    ) -> list[ParsedSegment]:
        """Parse a supported file."""
        return self.get_parser(path).parse(path, cancel_callback=cancel_callback)


class DocumentChunker:
    """Split parsed segments into retrieval-friendly chunks."""

    def __init__(self, settings: Settings | None = None) -> None:
        """Initialize chunker."""
        self.settings = settings or get_settings()

    def chunk(self, segments: list[ParsedSegment]) -> list[TextChunk]:
        """Chunk parsed segments while preserving table/code units."""
        chunks: list[TextChunk] = []
        for segment in segments:
            if segment.content_type in {"table", "code"}:
                chunks.extend(self._chunk_large_preserved_segment(segment, len(chunks)))
            else:
                chunks.extend(self._chunk_text_segment(segment, len(chunks)))
        for index, chunk in enumerate(chunks):
            chunk.chunk_index = index
        return chunks

    def _chunk_large_preserved_segment(
        self,
        segment: ParsedSegment,
        start_index: int,
    ) -> list[TextChunk]:
        """Keep table/code segments intact unless they exceed the chunk budget."""
        if estimate_tokens(segment.content) <= self.settings.chunk_size * 2:
            return [
                TextChunk(
                    content=segment.content,
                    content_type=segment.content_type,
                    chunk_index=start_index,
                    page_number=segment.page_number,
                    metadata=segment.metadata,
                )
            ]
        return self._chunk_text_segment(segment, start_index)

    def _chunk_text_segment(self, segment: ParsedSegment, start_index: int) -> list[TextChunk]:
        """Split a normal text segment by approximate tokens with overlap."""
        tokens = tokenize_for_chunking(segment.content)
        if not tokens:
            return []
        size = max(64, self.settings.chunk_size)
        overlap = min(max(0, self.settings.chunk_overlap), size // 2)
        step = max(1, size - overlap)
        chunks: list[TextChunk] = []
        for offset in range(0, len(tokens), step):
            token_slice = tokens[offset : offset + size]
            if not token_slice:
                continue
            text = untokenize(token_slice).strip()
            if not text:
                continue
            chunks.append(
                TextChunk(
                    content=text,
                    content_type=segment.content_type,
                    chunk_index=start_index + len(chunks),
                    page_number=segment.page_number,
                    metadata={**segment.metadata, "token_offset": offset},
                )
            )
            if offset + size >= len(tokens):
                break
        return chunks


class MetadataExtractor:
    """Extract source metadata and evidence-backed sensor hints."""

    MANUFACTURER_PATTERNS = (
        r"(?:制造商|厂商|品牌|Manufacturer|Brand)\s*[:：]\s*([A-Za-z0-9\u4e00-\u9fff ._\-&()]+)",
    )
    MODEL_PATTERNS = (
        r"(?:型号|产品型号|Model|Part\s*Number|P/N)\s*[:：]\s*([A-Za-z0-9_\-./]+)",
        r"\b([A-Z]{1,6}[-_]?\d{2,6}[A-Z0-9\-_.]*)\b",
    )
    # Explicit "型号: XXX" / "Model: XXX" labels are high-confidence evidence.
    MODEL_LABEL_PATTERNS = (MODEL_PATTERNS[0],)
    # Label-free fallback: an alphanumeric designator such as LDR-100 or VLP-16.
    MODEL_FALLBACK_PATTERN = MODEL_PATTERNS[1]
    # Common standards, interfaces, and protocols the fallback must not mistake
    # for a sensor model (for example ISO9001, RS232, IP67, GB2312, USB3).
    NON_MODEL_PREFIXES = frozenset(
        {
            "iso", "iec", "ieee", "ansi", "din", "jis", "gb", "gbt", "en",
            "ul", "ce", "fcc", "rohs", "reach", "mil", "nema",
            "rs", "ip", "usb", "hdmi", "vga", "tcp", "udp", "i2c", "spi",
            "can", "uart", "pwm", "adc", "dac", "led", "lcd", "pcb",
        }
    )

    def extract(
        self,
        path: str | Path,
        segments: list[ParsedSegment] | None = None,
        tags: list[str] | None = None,
        notes: str | None = None,
    ) -> dict[str, Any]:
        """Extract metadata from file info and parsed text evidence."""
        file_info = get_file_info(path)
        text = "\n".join(segment.content for segment in (segments or [])[:5])
        metadata: dict[str, Any] = {
            "filename": file_info.path.name,
            "file_path": str(file_info.path.resolve()),
            "file_type": file_info.file_type,
            "size_bytes": file_info.size_bytes,
            "created_at": file_info.created_at,
            "modified_at": file_info.modified_at,
            "tags": ",".join(tags or []),
            "notes": notes,
            "manufacturer": first_regex_group(text, self.MANUFACTURER_PATTERNS),
            "sensor_model": self._extract_model(text),
        }
        return metadata

    def extract_hints_from_text(self, text: str) -> dict[str, str | None]:
        """Derive sensor model and manufacturer hints from text evidence only.

        Shared by import-time extraction and the metadata refresh tool so both
        paths apply identical rules without re-parsing the source file.
        """
        return {
            "manufacturer": first_regex_group(text, self.MANUFACTURER_PATTERNS),
            "sensor_model": self._extract_model(text),
        }

    def _extract_model(self, text: str) -> str | None:
        """Return the most reliable sensor model evidence from text.

        Explicit ``型号:``/``Model:`` labels take priority. Otherwise the
        label-free fallback returns the first alphanumeric designator whose
        alphabetic prefix is not a known standard, interface, or protocol, so
        tokens like ISO9001, RS232, IP67 or USB3 are never treated as models.
        """
        labelled = first_regex_group(text, self.MODEL_LABEL_PATTERNS)
        if labelled:
            return labelled
        for match in re.finditer(self.MODEL_FALLBACK_PATTERN, text):
            candidate = match.group(1).strip(" ;,，。")
            prefix = re.match(r"[A-Za-z]+", candidate)
            if prefix and prefix.group(0).lower() in self.NON_MODEL_PREFIXES:
                continue
            return candidate
        return None


def _table_to_markdown(table: list[list[Any]]) -> str:
    """Convert an extracted table to Markdown."""
    cleaned = [["" if cell is None else str(cell).replace("\n", " ").strip() for cell in row] for row in table]
    cleaned = [row for row in cleaned if any(cell for cell in row)]
    if not cleaned:
        return ""
    max_columns = max(len(row) for row in cleaned)
    rows = [row + [""] * (max_columns - len(row)) for row in cleaned]
    header = rows[0]
    separator = ["---"] * max_columns
    body = rows[1:]
    markdown_rows = [header, separator, *body]
    return "\n".join("| " + " | ".join(row) + " |" for row in markdown_rows)


def _check_cancelled(cancel_callback: CancelCallback | None) -> None:
    """Run an optional cancellation checkpoint."""
    if cancel_callback:
        cancel_callback()


def tokenize_for_chunking(text: str) -> list[str]:
    """Tokenize Chinese/English text approximately for chunking."""
    return re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9_./%+\-]+|[^\s]", text)


def untokenize(tokens: list[str]) -> str:
    """Rebuild readable text from approximate tokens."""
    output: list[str] = []
    for token in tokens:
        if not output:
            output.append(token)
            continue
        prev = output[-1]
        if re.match(r"[\u4e00-\u9fff]", token) or re.match(r"[\u4e00-\u9fff]", prev[-1]):
            output.append(token)
        elif re.match(r"[,.!?;:%)\]}]", token):
            output.append(token)
        elif prev.endswith(("(", "[", "{", "/", "-")):
            output.append(token)
        else:
            output.append(" " + token)
    return "".join(output)


def estimate_tokens(text: str) -> int:
    """Estimate token count for chunking and validation."""
    return len(tokenize_for_chunking(text))


def first_regex_group(text: str, patterns: tuple[str, ...]) -> str | None:
    """Return the first captured regex group found in text."""
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip(" ;,，。")
    return None


def metadata_to_json(metadata: dict[str, Any]) -> str:
    """Serialize metadata for SQLite storage."""
    return json.dumps(metadata, ensure_ascii=False, default=str)
